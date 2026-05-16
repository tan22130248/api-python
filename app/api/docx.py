from fastapi import APIRouter, HTTPException
from fastapi.responses import StreamingResponse
from pydantic import BaseModel, field_validator
from typing import List, Optional
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io
import logging
import urllib.request
import urllib.error

router = APIRouter(prefix="/api/docx", tags=["docx"])
logger = logging.getLogger(__name__)

# Define Pydantic models
class AnswerDTO(BaseModel):
    id: Optional[int] = None
    label: str
    content: str
    isCorrect: bool

class QuestionDTO(BaseModel):
    id: Optional[int] = None
    type: str  # MULTIPLE_CHOICE, AUDIO, MATCHING, FILL_IN_BLANK, ESSAY
    content: Optional[str] = None
    points: int = 0
    title: Optional[str] = None
    numberQuestions: Optional[int] = None
    answers: Optional[List[AnswerDTO]] = None
    audioUrl: Optional[str] = None
    imageUrl: Optional[str] = None
    transcript: Optional[str] = None
    orderIndex: Optional[int] = None
    
    # New fields for MATCHING questions
    matchingPairs: Optional[List[dict]] = None
    
    # New fields for FILL_IN_BLANK questions
    textWithBlanks: Optional[str] = None
    blanks: Optional[List[dict]] = None
    
    # New fields for ESSAY questions
    prompt: Optional[str] = None
    maxLength: Optional[int] = None
    rubric: Optional[str] = None
    
    # Custom validator to handle both string and enum values
    @field_validator('type', mode='before')
    @classmethod
    def convert_type_to_string(cls, v):
        if v is None:
            return 'MULTIPLE_CHOICE'
        if hasattr(v, 'value'):  # Enum case
            return v.value
        return str(v)

class CreateTestRequest(BaseModel):
    name: str
    subject: str
    grade: Optional[str] = None
    duration: int
    description: Optional[str] = None
    lessonContentName: Optional[str] = None
    includeAnswers: Optional[bool] = True
    questions: List[QuestionDTO]

@router.post("/generate-test")
async def generate_test_docx(request: CreateTestRequest):
    """
    Generate a DOCX file from test data
    """
    try:
        logger.info(f"Generating DOCX for test: {request.name}")
        
        # Create a new Document
        doc = Document()
        
        # Add title
        title = doc.add_heading(request.name, level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add lesson content name centered below the title (if provided)
        if getattr(request, 'lessonContentName', None):
            lesson_para = doc.add_paragraph()
            lesson_run = lesson_para.add_run(str(request.lessonContentName))
            lesson_run.bold = True
            lesson_run.font.size = Pt(14)
            lesson_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add metadata
        metadata = doc.add_paragraph()
        metadata_run = metadata.add_run(f"Môn: {request.subject}")
        metadata_run.bold = True
        if request.grade:
            metadata.add_run(f" | Lớp: {request.grade}")
        metadata.add_run(f"\nThời gian: {request.duration} phút\n")
        metadata.add_run(f"Tổng điểm: {sum(q.points for q in request.questions)}")
        
        if request.description:
            doc.add_paragraph(f"Mô tả: {request.description}")
        
        doc.add_paragraph()  # Add spacing
        
        # Add questions
        for idx, question in enumerate(request.questions, 1):
            if question.type == "MULTIPLE_CHOICE":
                _add_multiple_choice_question(doc, idx, question, request.includeAnswers)
            elif question.type == "AUDIO":
                _add_audio_question(doc, idx, question)
            elif question.type == "MATCHING":
                _add_matching_question(doc, idx, question, request.includeAnswers)
            elif question.type == "FILL_IN_BLANK":
                _add_fill_in_blank_question(doc, idx, question, request.includeAnswers)
            elif question.type == "ESSAY":
                _add_essay_question(doc, idx, question)
            
            doc.add_paragraph()  # Add spacing between questions
        
        # Save to bytes
        doc_io = io.BytesIO()
        doc.save(doc_io)
        doc_io.seek(0)
        
        logger.info(f"DOCX generated successfully for test: {request.name}")
        
        # Create safe filename (ASCII only, replace special chars)
        safe_name = request.name.encode('ascii', 'replace').decode('ascii') if request.name else "test"
        
        # Return as StreamingResponse with proper headers
        return StreamingResponse(
            iter([doc_io.getvalue()]),
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f"attachment; filename={safe_name}.docx"}
        )
        
    except Exception as e:
        logger.error(f"Error generating DOCX: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Error generating DOCX: {str(e)}")

def _add_multiple_choice_question(doc, question_num: int, question: QuestionDTO, include_answers: bool = True):
    """Add a multiple choice question to the document"""
    
    # Question header
    header = doc.add_heading(f"Câu {question_num}: {question.title or 'Trắc nghiệm'}", level=2)
    
    # Question content - handle null/undefined
    content = question.content if question.content else ""
    content_para = doc.add_paragraph(f"Nội dung: {content} (Điểm: {question.points})")
    content_para.paragraph_format.left_indent = Inches(0.25)

    if question.imageUrl:
        doc.add_paragraph("Hình ảnh:")
        _add_image_from_url(doc, question.imageUrl)
    
    # Answers
    doc.add_paragraph("Các đáp án:")
    if question.answers:
        for answer in question.answers:
            answer_content = answer.content if answer.content else ""
            answer_text = f"{answer.label}. {answer_content}"
            if include_answers and answer.isCorrect:
                answer_text += " ✓ (Đáp án đúng)"
            answer_para = doc.add_paragraph(answer_text)
            answer_para.paragraph_format.left_indent = Inches(0.5)


def _add_image_from_url(doc, image_url: str):
    try:
        with urllib.request.urlopen(image_url, timeout=15) as response:
            image_bytes = response.read()
            image_stream = io.BytesIO(image_bytes)
            image_para = doc.add_paragraph()
            run = image_para.add_run()
            run.add_picture(image_stream, width=Inches(5.5))
            image_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    except Exception as e:
        logger.warning(f"Cannot download image from URL '{image_url}': {e}")
        fallback_para = doc.add_paragraph(f"Hình ảnh: {image_url}")
        fallback_para.paragraph_format.left_indent = Inches(0.25)


def _add_audio_question(doc, question_num: int, question: QuestionDTO):
    """Add an audio question to the document"""
    
    # Question header
    header = doc.add_heading(f"Câu {question_num}: Câu hỏi âm thanh", level=2)
    
    # Question content - handle null/undefined
    content = question.content if question.content else ""
    content_para = doc.add_paragraph(f"Nội dung: {content} (Điểm: {question.points})")
    content_para.paragraph_format.left_indent = Inches(0.25)
    
    # Audio info
    # Không xuất đường dẫn audio vào DOCX nếu chỉ cần nội dung câu hỏi âm thanh
    
    if question.imageUrl:
        doc.add_paragraph("Hình ảnh:")
        _add_image_from_url(doc, question.imageUrl)


def _add_matching_question(doc, question_num: int, question: QuestionDTO, include_answers: bool = True):
    """Add a matching question to the document"""
    
    # Question header
    header = doc.add_heading(f"Câu {question_num}: Nối các cặp từ", level=2)
    
    # Question content
    content = question.content if question.content else ""
    content_para = doc.add_paragraph(f"Nội dung: {content} (Điểm: {question.points})")
    content_para.paragraph_format.left_indent = Inches(0.25)
    
    # Create table for matching pairs
    if question.matchingPairs:
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Light Grid Accent 1'
        
        # Header row - swap columns if no answers
        hdr_cells = table.rows[0].cells
        if include_answers:
            hdr_cells[0].text = 'Cột trái'
            hdr_cells[1].text = 'Cột phải'
        else:
            hdr_cells[0].text = 'Cột phải'
            hdr_cells[1].text = 'Cột trái'
        
        # Add matching pairs - swap columns if no answers
        for pair in question.matchingPairs:
            row_cells = table.add_row().cells
            if include_answers:
                row_cells[0].text = pair.get('left', '')
                row_cells[1].text = pair.get('right', '')
            else:
                # For student version: show right column first, left column second
                row_cells[0].text = pair.get('right', '')
                row_cells[1].text = pair.get('left', '')
                
                # Set black color for text
                for cell in row_cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.color.rgb = RGBColor(0, 0, 0)  # Black color
        
        # Add empty row with "..." for student to fill answers if no answers
        if not include_answers:
            empty_row = table.add_row().cells
            empty_row[0].text = '...'
            empty_row[1].text = '...'
            # Set black color for the dots
            for cell in empty_row:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.color.rgb = RGBColor(0, 0, 0)


def _add_fill_in_blank_question(doc, question_num: int, question: QuestionDTO, include_answers: bool = True):
    """Add a fill-in-the-blank question to the document"""
    
    # Question header
    header = doc.add_heading(f"Câu {question_num}: Điền từ vào chỗ trống", level=2)
    
    # Question content
    content = question.content if question.content else ""
    content_para = doc.add_paragraph(f"Nội dung: {content} (Điểm: {question.points})")
    content_para.paragraph_format.left_indent = Inches(0.25)
    
    # Display text with blanks
    if question.textWithBlanks:
        # Replace blank markers with underscores
        display_text = question.textWithBlanks
        if question.blanks:
            for blank in question.blanks:
                blank_id = blank.get('id', '')
                display_text = display_text.replace(f'[BLANK_{blank_id}]', '___________')
        
        text_para = doc.add_paragraph(f"Đoạn văn: {display_text}")
        text_para.paragraph_format.left_indent = Inches(0.25)
        
        # Show answers if requested
        if include_answers and question.blanks:
            doc.add_paragraph("Đáp án:")
            for blank in question.blanks:
                answer_text = f"Chỗ trống {blank.get('id', '')}: {blank.get('correctAnswer', '')}"
                answer_para = doc.add_paragraph(answer_text)
                answer_para.paragraph_format.left_indent = Inches(0.5)


def _add_essay_question(doc, question_num: int, question: QuestionDTO):
    """Add an essay question to the document"""
    
    # Question header
    header = doc.add_heading(f"Câu {question_num}: Viết bài", level=2)
    
    # Question content
    content = question.content if question.content else ""
    content_para = doc.add_paragraph(f"Nội dung: {content} (Điểm: {question.points})")
    content_para.paragraph_format.left_indent = Inches(0.25)
    
    # Essay prompt
    if question.prompt:
        prompt_para = doc.add_paragraph(f"Yêu cầu: {question.prompt}")
        prompt_para.paragraph_format.left_indent = Inches(0.25)
    
    # Add space for answer
    doc.add_paragraph("Trả lời:")
    for i in range(8):  # Add more lines for essay
        doc.add_paragraph("_" * 80)
    
    # Rubric if available
    if question.rubric:
        rubric_para = doc.add_paragraph(f"Tiêu chí đánh giá: {question.rubric}")
        rubric_para.paragraph_format.left_indent = Inches(0.25)
        rubric_para.runs[0].italic = True

@router.post("/generate-test/stream")
async def generate_test_docx_stream(request: CreateTestRequest):
    """
    Generate and stream a DOCX file
    """
    try:
        # Generate DOCX bytes
        doc = Document()
        
        # Add title
        title = doc.add_heading(request.name, level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add metadata
        metadata = doc.add_paragraph()
        metadata_run = metadata.add_run(f"Môn: {request.subject}")
        metadata_run.bold = True
        if request.grade:
            metadata.add_run(f" | Lớp: {request.grade}")
        metadata.add_run(f"\nThời gian: {request.duration} phút\n")
        metadata.add_run(f"Tổng điểm: {sum(q.points for q in request.questions)}")
        
        if request.description:
            doc.add_paragraph(f"Mô tả: {request.description}")
        
        doc.add_paragraph()
        
        # Add questions
        for idx, question in enumerate(request.questions, 1):
            if question.type == "MULTIPLE_CHOICE":
                _add_multiple_choice_question(doc, idx, question, request.includeAnswers)
            elif question.type == "AUDIO":
                _add_audio_question(doc, idx, question)
            doc.add_paragraph()
        
        # Save to bytes
        doc_io = io.BytesIO()
        doc.save(doc_io)
        doc_io.seek(0)
        docx_bytes = doc_io.getvalue()
        
        # Return as base64 encoded string (not latin-1 decode!)
        import base64
        b64_data = base64.b64encode(docx_bytes).decode('utf-8')
        
        return {
            "success": True,
            "message": "DOCX generated successfully",
            "data": b64_data
        }
    except Exception as e:
        logger.error(f"Error generating DOCX stream: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Error: {str(e)}")
